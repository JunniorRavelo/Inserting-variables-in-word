from docx import Document  # Importa la clase Document desde el módulo docx
from docx.shared import Pt  # Importa la clase Pt desde el módulo shared de docx

def obtener_estilo(parrafo):
    estilo = {}  # Crea un diccionario para almacenar el estilo del párrafo
    if parrafo.style.name:  # Verifica si el párrafo tiene un estilo definido
        estilo['name'] = parrafo.style.name  # Almacena el nombre del estilo en el diccionario
    if parrafo.style.font.bold:  # Verifica si el texto del párrafo está en negrita
        estilo['bold'] = True  # Almacena True si el texto está en negrita en el diccionario
    if parrafo.style.font.italic:  # Verifica si el texto del párrafo está en cursiva
        estilo['italic'] = True  # Almacena True si el texto está en cursiva en el diccionario
    if parrafo.style.font.underline:  # Verifica si el texto del párrafo está subrayado
        estilo['underline'] = True  # Almacena True si el texto está subrayado en el diccionario
    if parrafo.style.font.color.rgb:  # Verifica si el texto del párrafo tiene un color definido
        estilo['color'] = parrafo.style.font.color.rgb  # Almacena el color del texto en el diccionario
    if parrafo.style.font.size:  # Verifica si el texto del párrafo tiene un tamaño definido
        estilo['size'] = parrafo.style.font.size  # Almacena el tamaño del texto en el diccionario
    return estilo  # Devuelve el diccionario con el estilo del párrafo


def aplicar_estilo(texto, estilo):
    # Verifica si el nombre del estilo está presente en el diccionario
    if 'name' in estilo:
        texto.style = estilo['name']  # Aplica el nombre del estilo al texto
    
    # Verifica si el estilo de negrita está presente en el diccionario
    if 'bold' in estilo:
        texto.bold = estilo['bold']  # Aplica el estilo de negrita al texto
    
    # Verifica si el estilo de cursiva está presente en el diccionario
    if 'italic' in estilo:
        texto.italic = estilo['italic']  # Aplica el estilo de cursiva al texto
    
    # Verifica si el estilo de subrayado está presente en el diccionario
    if 'underline' in estilo:
        texto.underline = estilo['underline']  # Aplica el estilo de subrayado al texto
    
    # Verifica si el color del texto está presente en el diccionario
    if 'color' in estilo:
        texto.font.color.rgb = estilo['color']  # Aplica el color al texto
    
    # Verifica si el tamaño del texto está presente en el diccionario
    if 'size' in estilo:
        texto.font.size = estilo['size']  # Aplica el tamaño al texto


def remplazar_variables(documento, variables):
    # Recorre todos los párrafos en el documento
    for p in documento.paragraphs:
        # Por cada variable y valor en el diccionario de variables
        for var, valor in variables.items():
            # Verifica si la variable está presente en el texto del párrafo
            if var in p.text:
                partes = p.runs  # Obtiene todas las partes del párrafo
                for parte in partes:
                    estilo = obtener_estilo(parte)  # Obtiene el estilo de la parte
                    texto_reemplazo = parte.text.replace(var, valor)  # Reemplaza la variable por su valor en el texto de la parte
                    parte.clear()  # Borra el contenido de la parte
                    parte.text = texto_reemplazo  # Establece el texto reemplazado en la parte existente
                    aplicar_estilo(parte, estilo)  # Aplica el estilo a la parte

    # Recorre todas las tablas en el documento
    for table in documento.tables:
        # Recorre todas las filas en la tabla
        for row in table.rows:
            # Recorre todas las celdas en la fila
            for cell in row.cells:
                # Por cada variable y valor en el diccionario de variables
                for var, valor in variables.items():
                    # Verifica si la variable está presente en el texto de la celda
                    if var in cell.text:
                        partes = cell.paragraphs[0].runs  # Obtiene todas las partes del texto de la celda
                        for parte in partes:
                            estilo = obtener_estilo(parte)  # Obtiene el estilo de la parte
                            texto_reemplazo = parte.text.replace(var, valor)  # Reemplaza la variable por su valor en el texto de la parte
                            parte.clear()  # Borra el contenido de la parte
                            parte.text = texto_reemplazo  # Establece el texto reemplazado en la parte existente
                            aplicar_estilo(parte, estilo)  # Aplica el estilo a la parte


# Ruta del archivo de Word
ruta_archivo = "spanish.docx"

# Cargar el documento de Word
doc = Document(ruta_archivo)

# Variables y sus valores
variables = {
    "[var1]": "Viaje",
    "[var2]": "Conejolandia",
    "[var3]": "Saltarín",
    "[var4]": "Fantasía",
    "[var5]": "Aladino"
}

# Reemplazar las variables en el documento
remplazar_variables(doc, variables)

# Guardar el documento modificado
doc.save("successfully_edited.docx")
